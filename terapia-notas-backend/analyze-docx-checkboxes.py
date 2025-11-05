#!/usr/bin/env python3
"""
Script para analizar documentos Word generados y verificar que los checkboxes
y labels de goals se estén reemplazando correctamente.
"""

import sys
import zipfile
import os
from io import BytesIO
import urllib.request
import urllib.parse
import json
import re

# Configuración
BASE_URL = "http://localhost:3002"
API_ENDPOINT = "/api/v1/notes/generate-group-week"

def extract_text_from_docx(docx_path):
    """Extrae texto de un archivo DOCX"""
    try:
        import docx
        doc = docx.Document(docx_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        full_text = '\n'.join(text)
        # Buscar también en tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text.append(para.text)
        return '\n'.join(text)
    except ImportError:
        print("[WARNING] python-docx no esta instalado. Instalalo con: pip install python-docx")
        print("   Intentando extraer texto usando zipfile...")
        try:
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                # Leer el contenido XML del documento
                xml_content = zip_ref.read('word/document.xml').decode('utf-8')
                # Extraer texto usando regex simple (mejorado)
                # Buscar en todos los elementos de texto, incluyendo espacios
                text_parts = []
                # Buscar patrones de texto más amplios
                for match in re.finditer(r'<w:t[^>]*>([^<]+)</w:t>', xml_content):
                    text_parts.append(match.group(1))
                # También buscar en elementos con espacios
                for match in re.finditer(r'<w:t[^>]*xml:space="preserve"[^>]*>([^<]+)</w:t>', xml_content):
                    text_parts.append(match.group(1))
                return '\n'.join(text_parts)
        except Exception as e:
            print(f"[ERROR] No se pudo extraer texto: {str(e)}")
            return None

def analyze_goals_in_text(text):
    """Analiza el texto para verificar checkboxes y labels de goals"""
    if not text:
        return None
    
    results = {
        'goal1': {'checkbox': None, 'label': None, 'text': None},
        'goal2': {'checkbox': None, 'label': None, 'text': None},
        'goal3': {'checkbox': None, 'label': None, 'text': None},
        'goal4': {'checkbox': None, 'label': None, 'text': None},
    }
    
    # Buscar patrones de goals (usar codigos Unicode para los checkboxes)
    patterns = [
        (r'([☒☐])\s*GOAL#1[:]\s*(.+)', 'goal1'),
        (r'([☒☐])\s*GOAL#2[:]\s*(.+)', 'goal2'),
        (r'([☒☐])\s*GOAL#3[:]\s*(.+)', 'goal3'),
        (r'([☒☐])\s*GOAL#4[:]\s*(.+)', 'goal4'),
        # Tambien buscar sin espacio
        (r'([☒☐])GOAL#1[:]\s*(.+)', 'goal1'),
        (r'([☒☐])GOAL#2[:]\s*(.+)', 'goal2'),
        (r'([☒☐])GOAL#3[:]\s*(.+)', 'goal3'),
        (r'([☒☐])GOAL#4[:]\s*(.+)', 'goal4'),
    ]
    
    for pattern, goal_key in patterns:
        match = re.search(pattern, text)
        if match:
            checkbox = match.group(1)
            goal_text = match.group(2).strip()
            results[goal_key]['checkbox'] = checkbox
            results[goal_key]['label'] = f'GOAL#{goal_key[-1]}'
            results[goal_key]['text'] = goal_text[:100]  # Primeros 100 caracteres
    
    # Si no encontramos con checkbox, buscar sin checkbox
    if not any(r['checkbox'] for r in results.values()):
        patterns_no_checkbox = [
            (r'GOAL#1:\s*(.+)', 'goal1'),
            (r'GOAL#2:\s*(.+)', 'goal2'),
            (r'GOAL#3:\s*(.+)', 'goal3'),
            (r'GOAL#4:\s*(.+)', 'goal4'),
        ]
        
        for pattern, goal_key in patterns_no_checkbox:
            match = re.search(pattern, text)
            if match:
                goal_text = match.group(1).strip()
                results[goal_key]['checkbox'] = 'NO ENCONTRADO'
                results[goal_key]['label'] = f'GOAL#{goal_key[-1]}'
                results[goal_key]['text'] = goal_text[:100]
    
    return results

def analyze_client_response_labels(text):
    """Analiza los labels de client response"""
    labels = {
        'group1': None,
        'group2': None,
        'group3': None,
        'group4': None,
    }
    
    # Patrones más flexibles para capturar diferentes formatos
    patterns = [
        (r'Group 1[^:]*Client Response[^:]*[:\s]*(.+)', 'group1'),
        (r'Group 2[^:]*Client Response[^:]*[:\s]*(.+)', 'group2'),
        (r'Group 3[^:]*Client Response[^:]*[:\s]*(.+)', 'group3'),
        (r'Group 4[^:]*Client Response[^:]*[:\s]*(.+)', 'group4'),
        # Buscar también con Goal# en el label
        (r'Group 1[^:]*Client Response.*Goal#\d+[^:]*[:\s]*(.+)', 'group1'),
        (r'Group 2[^:]*Client Response.*Goal#\d+[^:]*[:\s]*(.+)', 'group2'),
        (r'Group 3[^:]*Client Response.*Goal#\d+[^:]*[:\s]*(.+)', 'group3'),
        (r'Group 4[^:]*Client Response.*Goal#\d+[^:]*[:\s]*(.+)', 'group4'),
    ]
    
    for pattern, group_key in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # Capturar el texto completo del match
            full_match = match.group(0)
            if full_match:
                # Buscar si tiene Goal# en el label
                goal_match = re.search(r'Goal#(\d+)', full_match)
                label_text = full_match[:120]  # Primeros 120 caracteres
                if goal_match:
                    label_text = f"Group {group_key[-1]}: Client Response(Goal#{goal_match.group(1)}/Obj{goal_match.group(1)}A): {full_match[:80]}..."
                labels[group_key] = label_text
                break  # Solo tomar el primer match
    
    return labels

def analyze_document(docx_path):
    """Analiza un documento Word completo"""
    print(f"\n{'='*60}")
    print(f"Analizando: {os.path.basename(docx_path)}")
    print(f"{'='*60}")
    
    text = extract_text_from_docx(docx_path)
    if not text:
        print("[ERROR] No se pudo extraer texto del documento")
        return None
    
    # Mostrar preview del texto extraído (primeros 500 caracteres)
    print(f"\n[PREVIEW] Primeros 500 caracteres del texto extraido:")
    if text:
        # Reemplazar caracteres problemáticos para evitar errores de encoding
        safe_text = text[:500].encode('ascii', 'replace').decode('ascii')
        print(safe_text)
    else:
        print("No hay texto")
    
    # Analizar goals
    goals = analyze_goals_in_text(text)
    if goals:
        print("\n[GOALS] Analisis de checkboxes y labels:")
        goals_found = False
        for goal_key, goal_data in goals.items():
            checkbox = goal_data['checkbox'] or 'NO ENCONTRADO'
            label = goal_data['label'] or 'NO ENCONTRADO'
            text_preview = goal_data['text'] or 'NO ENCONTRADO'
            
            if checkbox != 'NO ENCONTRADO':
                goals_found = True
                # Usar caracteres ASCII para evitar problemas de encoding
                if checkbox == '☒':
                    checkbox_display = '[X]'
                elif checkbox == '☐':
                    checkbox_display = '[ ]'
                else:
                    checkbox_display = checkbox
                print(f"  {goal_key.upper()}: {checkbox_display} {label}")
                if text_preview != 'NO ENCONTRADO':
                    print(f"    Texto: {text_preview[:80]}...")
        
        if not goals_found:
            print("  [WARNING] No se encontraron goals con checkboxes en el texto")
            # Buscar cualquier referencia a GOAL en el texto
            goal_matches = re.findall(r'GOAL#\d+', text, re.IGNORECASE)
            if goal_matches:
                print(f"  [INFO] Se encontraron referencias a: {set(goal_matches)}")
    
    # Analizar client response labels
    labels = analyze_client_response_labels(text)
    if labels:
        print("\n[CLIENT RESPONSE LABELS] Analisis de labels:")
        labels_found = False
        for group_key, label_text in labels.items():
            if label_text:
                labels_found = True
                # Verificar si tiene Goal# en el label
                if 'Goal#' in label_text:
                    print(f"  {group_key.upper()}: {label_text[:100]}... [CON Goal#]")
                else:
                    print(f"  {group_key.upper()}: {label_text[:100]}... [SIN Goal#]")
        
        if not labels_found:
            print("  [WARNING] No se encontraron labels de client response")
            # Buscar cualquier referencia a "Client Response" en el texto
            response_matches = re.findall(r'Group\s+\d+[^:]*Client\s+Response[^:]*', text, re.IGNORECASE)
            if response_matches:
                print(f"  [INFO] Se encontraron referencias a Client Response:")
                for match in response_matches[:4]:
                    print(f"    - {match[:80]}...")
    
    return {
        'goals': goals,
        'labels': labels,
    }

def main():
    """Funcion principal"""
    if len(sys.argv) < 3:
        print("Uso: python analyze-docx-checkboxes.py <groupId> <weekId>")
        print("Ejemplo: python analyze-docx-checkboxes.py 81a01050-1b2c-4da4-9db3-ce4debb41fd0 0d9cf86d-1191-49b7-9bb0-5abe4a5d23e5")
        return 1
    
    group_id = sys.argv[1]
    week_id = sys.argv[2]
    
    print("="*60)
    print("ANALISIS DE DOCUMENTOS WORD GENERADOS")
    print("="*60)
    
    # Generar notas
    print(f"\n[1] Generando notas para groupId={group_id}, weekId={week_id}...")
    url = f"{BASE_URL}{API_ENDPOINT}"
    payload = {
        "groupId": group_id,
        "weekId": week_id
    }
    
    try:
        data = json.dumps(payload).encode('utf-8')
        req = urllib.request.Request(url, data=data, method='POST')
        req.add_header('Content-Type', 'application/json')
        
        print(f"    Peticion POST a: {url}")
        
        with urllib.request.urlopen(req, timeout=120) as response:
            status_code = response.getcode()
            if status_code != 200:
                print(f"[ERROR] Status Code: {status_code}")
                return 1
            
            content = response.read()
            zip_file = zipfile.ZipFile(BytesIO(content))
            file_list = zip_file.namelist()
            
            print(f"    [OK] ZIP generado con {len(file_list)} archivos")
            
            # Extraer y analizar el primer documento de cada dia
            print(f"\n[2] Analizando documentos generados...")
            
            # Agrupar por dia (1027, 1028, etc.)
            files_by_day = {}
            for filename in file_list:
                if filename.endswith('.docx'):
                    # Extraer fecha del nombre (ej: 1027.docx)
                    date_match = re.search(r'(\d{4})\.docx', filename)
                    if date_match:
                        date_str = date_match.group(1)
                        if date_str not in files_by_day:
                            files_by_day[date_str] = []
                        files_by_day[date_str].append(filename)
            
            # Analizar un documento de cada dia
            for date_str in sorted(files_by_day.keys())[:5]:  # Solo primeros 5 dias
                files = files_by_day[date_str]
                if files:
                    # Tomar el primer archivo del dia
                    filename = files[0]
                    
                    # Extraer a archivo temporal (limpiar nombre de archivo)
                    safe_filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
                    safe_filename = safe_filename.replace('/', '_').replace('\\', '_')
                    temp_path = f"temp_{date_str}_{safe_filename[-50:]}"  # Limitar longitud
                    with open(temp_path, 'wb') as f:
                        f.write(zip_file.read(filename))
                    
                    # Analizar
                    result = analyze_document(temp_path)
                    
                    # Limpiar
                    try:
                        os.remove(temp_path)
                    except:
                        pass
    
    except Exception as e:
        print(f"[ERROR] Error durante la generacion o analisis: {str(e)}")
        import traceback
        traceback.print_exc()
        return 1
    
    print("\n" + "="*60)
    print("ANALISIS COMPLETADO")
    print("="*60)
    
    return 0

if __name__ == "__main__":
    exit_code = main()
    sys.exit(exit_code)

